from io import BytesIO
from pathlib import Path

import pytest
from docx import Document

from app.exporters.docx_exporter import DocxExporter


def test_docx_exporter_rejects_empty_content() -> None:
    with pytest.raises(ValueError, match="must not be empty"):
        DocxExporter().export(" \n\t ")


def test_docx_exporter_returns_docx_bytes() -> None:
    docx = DocxExporter().export("# Jane Example\n\nBuilds local tools.")

    assert isinstance(docx, bytes)
    assert docx.startswith(b"PK")


def test_docx_can_be_opened_with_python_docx() -> None:
    docx = DocxExporter().export("# Jane Example\n\nBuilds local tools.")

    document = Document(BytesIO(docx))

    assert [paragraph.text for paragraph in document.paragraphs] == [
        "Jane Example",
        "Builds local tools.",
    ]


def test_docx_exporter_renders_headings_bullet_lists_and_paragraphs() -> None:
    markdown = (
        "# Jane Example\n\n## Skills\n\n- Python\n- FastAPI\n\nBuilds local tools."
    )

    document = Document(BytesIO(DocxExporter().export(markdown)))
    paragraph_texts = [paragraph.text for paragraph in document.paragraphs]

    assert "Jane Example" in paragraph_texts
    assert "Skills" in paragraph_texts
    assert "Python" in paragraph_texts
    assert "FastAPI" in paragraph_texts
    assert "Builds local tools." in paragraph_texts


def test_docx_exporter_treats_raw_html_and_script_like_input_as_text() -> None:
    markdown = "# <script>alert('x')</script>\n\n<script src=\"https://example.invalid/x.js\"></script>"

    document = Document(BytesIO(DocxExporter().export(markdown, title="<Tailored>")))
    paragraph_texts = [paragraph.text for paragraph in document.paragraphs]

    assert "<script>alert('x')</script>" in paragraph_texts
    assert '<script src="https://example.invalid/x.js"></script>' in paragraph_texts


def test_docx_exporter_does_not_write_files_directly(tmp_path: Path) -> None:
    before = set(tmp_path.iterdir())

    DocxExporter().export("# Tailored CV\n\n- Python")

    assert set(tmp_path.iterdir()) == before
