from __future__ import annotations

from io import BytesIO
from typing import Any

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

BRAND_BLUE = RGBColor(11, 74, 111)


class DocxExporter:
    """Render styled resume content as DOCX bytes."""

    def export(self, markdown: str, title: str = "Tailored CV") -> bytes:
        if not markdown.strip():
            raise ValueError("DOCX export content must not be empty.")
        return self.export_content(_markdown_to_content(markdown, title), title=title)

    def export_content(
        self, content: dict[str, Any], title: str = "Tailored CV"
    ) -> bytes:
        if not content:
            raise ValueError("DOCX export content must not be empty.")
        document = Document()
        document.core_properties.title = title
        _configure_document(document)
        _render_content(document, content)
        buffer = BytesIO()
        document.save(buffer)
        return buffer.getvalue()


def _configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Pt(36)
    section.bottom_margin = Pt(36)
    section.left_margin = Pt(40)
    section.right_margin = Pt(40)
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9.5)


def _render_content(document: Document, content: dict[str, Any]) -> None:
    sections = content.get("sections", {})
    header = sections.get("header", {})
    name = " ".join(
        part
        for part in [header.get("first_name", ""), header.get("surname", "")]
        if part
    ).strip() or content.get("name", "Resume")
    p = document.add_paragraph()
    run = p.add_run(_clean_text(name).upper())
    run.font.color.rgb = BRAND_BLUE
    run.font.size = Pt(16)
    run.bold = False
    if content.get("target_role"):
        p = document.add_paragraph()
        run = p.add_run(_clean_text(str(content["target_role"])).upper())
        run.font.color.rgb = BRAND_BLUE
        run.font.size = Pt(13)
    contact = " • ".join(
        _clean_text(str(part))
        for part in [
            header.get("phone", ""),
            header.get("email", ""),
            header.get("linkedin_url", ""),
            header.get("github_url", ""),
            header.get("location", ""),
            header.get("extra_text", ""),
        ]
        if part
    )
    if contact:
        document.add_paragraph(contact)
    summary = sections.get("summary", {}).get("text", "").strip()
    if summary:
        document.add_paragraph(_clean_text(summary))
    skills = sections.get("skills", {})
    if skills.get("hard") or skills.get("soft"):
        _add_section_heading(document, "Skills")
        if skills.get("hard"):
            _add_labelled_paragraph(document, "Hard Skills", skills["hard"])
        if skills.get("soft"):
            _add_labelled_paragraph(document, "Soft Skills", skills["soft"])
    _add_experience(
        document, sections.get("work_experience", []), "Professional Experience"
    )
    _add_education(document, sections.get("education", []))
    _add_rows(document, sections.get("languages", []), "Languages", _language_line)
    _add_rows(
        document, sections.get("certificates", []), "Certificates", _certificate_line
    )
    _add_rows(document, sections.get("references", []), "References", _reference_line)


def _add_section_heading(document: Document, title: str) -> None:
    p = document.add_paragraph()
    _set_paragraph_border(p)
    run = p.add_run(title.upper())
    run.font.color.rgb = BRAND_BLUE
    run.font.size = Pt(12)


def _add_labelled_paragraph(document: Document, label: str, text: str) -> None:
    p = document.add_paragraph()
    label_run = p.add_run(f"{label}: ")
    label_run.bold = True
    p.add_run(_clean_text(text))


def _add_experience(document: Document, rows: list[dict[str, Any]], title: str) -> None:
    visible = [
        row
        for row in rows
        if row.get("role_title") or row.get("organisation") or row.get("content")
    ]
    if not visible:
        return
    _add_section_heading(document, title)
    for row in visible:
        heading = " at ".join(
            part for part in [row.get("role_title"), row.get("organisation")] if part
        )
        if heading:
            p = document.add_paragraph()
            run = p.add_run(_clean_text(heading))
            run.font.color.rgb = BRAND_BLUE
            run.font.size = Pt(10.5)
        period = _period(row)
        if period:
            p = document.add_paragraph()
            p.add_run(_clean_text(period)).bold = True
        if row.get("optional_extra_enabled") and row.get("optional_extra_text"):
            document.add_paragraph(_clean_text(row["optional_extra_text"]))
        _add_bullets(document, row.get("content", ""))


def _add_education(document: Document, rows: list[dict[str, Any]]) -> None:
    visible = [
        row
        for row in rows
        if row.get("organisation") or row.get("role_title") or row.get("content")
    ]
    if not visible:
        return
    _add_section_heading(document, "Education")
    for row in visible:
        heading = " — ".join(
            part for part in [row.get("organisation"), row.get("role_title")] if part
        )
        if heading:
            document.add_paragraph(_clean_text(heading))
        period = _period(row)
        if period:
            document.add_paragraph(_clean_text(period))
        _add_bullets(document, row.get("content", ""))


def _add_rows(
    document: Document, rows: list[dict[str, Any]], title: str, formatter
) -> None:  # type: ignore[no-untyped-def]
    rendered = [_clean_text(formatter(row)) for row in rows]
    rendered = [item for item in rendered if item and item != "()"]
    if not rendered:
        return
    _add_section_heading(document, title)
    for item in rendered:
        document.add_paragraph(item, style="List Bullet")


def _add_bullets(document: Document, text: str) -> None:
    for raw in text.splitlines():
        bullet = _clean_text(raw.strip().lstrip("-• ").strip())
        if bullet:
            document.add_paragraph(bullet, style="List Bullet")


def _set_paragraph_border(paragraph) -> None:  # type: ignore[no-untyped-def]
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    for edge in ["top", "bottom"]:
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "2")
        element.set(qn("w:color"), "111827")
        borders.append(element)
    p_pr.append(borders)


def _period(row: dict[str, Any]) -> str:
    start = row.get("start_date", "")
    end = "Current" if row.get("is_current") else row.get("end_date", "")
    return " – ".join(part for part in [start, end] if part)


def _language_line(row: dict[str, Any]) -> str:
    language = row.get("language", row.get("title", ""))
    level = row.get("level", row.get("subtitle", ""))
    return f"{language} ({level})".strip()


def _certificate_line(row: dict[str, Any]) -> str:
    text = " | ".join(
        part
        for part in [
            row.get("certificate_name", row.get("title", "")),
            row.get("issue_year", ""),
        ]
        if part
    )
    return f"{text} — {row['certificate_url']}" if row.get("certificate_url") else text


def _reference_line(row: dict[str, Any]) -> str:
    return " — ".join(
        part
        for part in [
            row.get("name", row.get("title", "")),
            ", ".join(
                part
                for part in [row.get("role_title", ""), row.get("company", "")]
                if part
            ),
            " • ".join(
                part
                for part in [
                    row.get("phone", ""),
                    row.get("email", ""),
                    row.get("linkedin_url", ""),
                ]
                if part
            ),
        ]
        if part
    )


def _clean_text(value: str) -> str:
    return (
        value.replace("**", "")
        .replace("##", "")
        .replace("#", "")
        .replace("\u200b", "")
        .strip()
    )


def _markdown_to_content(markdown: str, title: str) -> dict[str, Any]:
    lines = [line.strip() for line in markdown.splitlines() if line.strip()]
    name = lines[0].lstrip("# ") if lines else title
    return {
        "name": title,
        "target_role": "",
        "sections": {"header": {"first_name": name}},
    }
